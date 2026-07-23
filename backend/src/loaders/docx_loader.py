from docx import Document
from langchain_core.documents import Document as LangChainDocument

from pathlib import Path
import uuid


def load_docx(file_path: str):
    """
    Extract text from DOCX file.

    Returns:
        List of LangChain Documents
    """

    doc = Document(file_path)

    text = []

    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text.append(paragraph.text)

    full_text = "\n".join(text)

    document_id = str(uuid.uuid4())

    document = LangChainDocument(
        page_content=full_text,
        metadata={
            "document_id": document_id,
            "source": "docx",
            "file": Path(file_path).name,
        }
    )

    return [document]